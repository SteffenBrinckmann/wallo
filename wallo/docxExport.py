"""DOCX export functionality for the Wallo application."""
import re
from pathlib import Path
from typing import Optional
from PySide6.QtCore import QObject
from PySide6.QtGui import QTextEdit, QFont, QColor
from PySide6.QtWidgets import QMessageBox
import pypandoc


class DocxExporter(QObject):
    """Handles DOCX export functionality with formatting preservation."""
    
    def __init__(self, parent: Optional[QObject] = None) -> None:
        super().__init__(parent)
    
    
    def exportToDocx(self, editor: QTextEdit, filename: str) -> None:
        """Export QTextEdit content to DOCX file with formatting preservation.
        
        Args:
            editor: The QTextEdit widget containing the content
            filename: The output filename for the DOCX file
        """
        try:
            # Try to use python-docx for better formatting preservation
            try:
                from docx import Document
                from docx.shared import RGBColor
                self.exportWithPythonDocx(editor, filename)
                self.showSuccessMessage(filename)
            except ImportError:
                # Fallback to pypandoc if python-docx is not available
                self.exportWithPypandoc(editor, filename)
                self.showSuccessMessage(filename, fallback=True)
                
        except Exception as e:
            self.showErrorMessage(str(e))
    
    
    def exportWithPythonDocx(self, editor: QTextEdit, filename: str) -> None:
        """Export content using python-docx library for better formatting preservation.
        
        Args:
            editor: The QTextEdit widget containing the content
            filename: The output filename for the DOCX file
        """
        from docx import Document
        from docx.shared import RGBColor
        
        document = Document()
        
        # Get the document content
        textDocument = editor.document()
        
        # Process each block (paragraph) in the document
        for blockNumber in range(textDocument.blockCount()):
            block = textDocument.findBlockByNumber(blockNumber)
            if not block.isValid():
                continue
                
            # Create a new paragraph for each block
            paragraph = document.add_paragraph()
            
            # Get the block iterator
            blockIterator = block.begin()
            
            # Process fragments within the block
            while not blockIterator.atEnd():
                fragment = blockIterator.fragment()
                if fragment.isValid():
                    fragmentText = fragment.text()
                    charFormat = fragment.charFormat()
                    
                    # Create a run for this fragment
                    run = paragraph.add_run(fragmentText)
                    
                    # Apply formatting based on character format
                    if charFormat.fontWeight() == QFont.Bold:
                        run.bold = True
                    if charFormat.fontItalic():
                        run.italic = True  
                    if charFormat.fontUnderline():
                        run.underline = True
                    
                    # Apply color if it's not default
                    color = charFormat.foreground().color()
                    if color.isValid() and color != QColor(0, 0, 0):  # Not default black
                        run.font.color.rgb = RGBColor(color.red(), color.green(), color.blue())
                
                blockIterator += 1
        
        document.save(filename)
    
    
    def exportWithPypandoc(self, editor: QTextEdit, filename: str) -> None:
        """Export content using pypandoc as fallback method.
        
        Args:
            editor: The QTextEdit widget containing the content
            filename: The output filename for the DOCX file
        """
        # Get HTML content from editor
        html = editor.toHtml()
        
        # Clean up HTML for better DOCX conversion
        html = self.cleanHtmlForDocx(html)
        
        # Ensure pandoc is available
        pypandoc.download_pandoc()
        
        # Convert HTML to DOCX with better formatting preservation
        extraArgs = ['--preserve-tabs']
        if Path('reference.docx').exists():
            extraArgs.append('--reference-doc=reference.docx')
        
        pypandoc.convert_text(
            html,
            'docx',
            format='html',
            outputfile=filename,
            extra_args=extraArgs
        )
    
    
    def cleanHtmlForDocx(self, html: str) -> str:
        """Clean HTML content for better DOCX conversion.
        
        Args:
            html: Raw HTML from QTextEdit
            
        Returns:
            Cleaned HTML suitable for DOCX conversion
        """
        # Convert Qt-specific color formatting to standard HTML
        html = re.sub(r'style="[^"]*color:\s*rgb\(([^)]+)\)[^"]*"', 
                     lambda m: f'style="color: rgb({m.group(1)})"', html)
        
        # Convert Qt font-weight to standard HTML
        html = re.sub(r'style="[^"]*font-weight:\s*([^;]+)[^"]*"', 
                     lambda m: f'style="font-weight: {m.group(1)}"', html)
        
        # Convert Qt font-style to standard HTML
        html = re.sub(r'style="[^"]*font-style:\s*([^;]+)[^"]*"', 
                     lambda m: f'style="font-style: {m.group(1)}"', html)
        
        # Convert Qt text-decoration to standard HTML
        html = re.sub(r'style="[^"]*text-decoration:\s*([^;]+)[^"]*"', 
                     lambda m: f'style="text-decoration: {m.group(1)}"', html)
        
        # Remove Qt-specific margin and padding that might interfere
        html = re.sub(r'margin-top:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-bottom:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-left:\s*[^;]+;?', '', html)
        html = re.sub(r'margin-right:\s*[^;]+;?', '', html)
        
        # Clean up empty style attributes
        html = re.sub(r'style=""', '', html)
        html = re.sub(r'style="[;\s]*"', '', html)
        
        # Ensure proper HTML structure
        if not html.startswith('<html'):
            html = f'<html><head><meta charset="UTF-8"></head><body>{html}</body></html>'
        
        return html
    
    
    def showSuccessMessage(self, filename: str, fallback: bool = False) -> None:
        """Show success message to user.
        
        Args:
            filename: The saved filename
            fallback: Whether pypandoc fallback was used
        """
        if fallback:
            message = f"Document saved successfully as {filename} (using pypandoc)"
        else:
            message = f"Document saved successfully as {filename}"
        
        QMessageBox.information(self.parent(), "Success", message)
    
    
    def showErrorMessage(self, error: str) -> None:
        """Show error message to user.
        
        Args:
            error: The error message to display
        """
        QMessageBox.critical(self.parent(), "Save Error", f"Failed to save document: {error}")