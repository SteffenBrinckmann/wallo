import sys, json
from pathlib import Path
from PySide6.QtWidgets import QApplication, QMainWindow, QTextEdit, QToolBar, QFileDialog, QMessageBox
from PySide6.QtGui import QIcon, QTextCursor, QTextCharFormat, QFont, QAction
from docx import Document
from bs4 import BeautifulSoup
from openai import OpenAI
from .fixedStrings import defaultConfiguration


class Wallo(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("WALLO - Writing Assistance by Large Language mOdel")
        self.editor = QTextEdit()
        self.setCentralWidget(self.editor)
        self._create_toolbar()
        self.configFile = Path.home()/'.wallo.json'
        if not self.configFile.is_file():
            with open(self.configFile, 'w', encoding='utf-8') as confFile:
                confFile.write(json.dumps(defaultConfiguration, indent=2))

    def _create_toolbar(self):
        toolbar = QToolBar("Formatting")
        self.addToolBar(toolbar)
        boldAction = QAction("Bold", self)           # Bold
        boldAction.triggered.connect(self.toggleBold)
        toolbar.addAction(boldAction)
        italicAction = QAction("Italic", self)        # Italic
        italicAction.triggered.connect(self.toggleItalic)
        toolbar.addAction(italicAction)
        underlineAction = QAction("Underline", self)         # Underline
        underlineAction.triggered.connect(self.toggleUnderline)
        toolbar.addAction(underlineAction)
        useLLM = QAction("Use LLM", self)         # use LLM
        useLLM.triggered.connect(self.addLLM)
        toolbar.addAction(useLLM)
        saveAction = QAction("Save as docx", self)        # Save as HTML
        saveAction.triggered.connect(self.saveDocx)
        toolbar.addAction(saveAction)

    def addLLM(self):
        cursor = self.editor.textCursor()
        if not cursor.hasSelection():
            QMessageBox.information(self, "Warning", "You have to select text for the tool to work")
            return
        conf = json.load(open(self.configFile, 'r', encoding='utf-8'))
        service = conf['services']['openAI']
        client = OpenAI(api_key=service['api'], base_url=service['url'])
        prompt = conf['prompts'][0]['user-prompt'] + '\n' + self.editor.textCursor().selectedText()
        response = client.chat.completions.create(model=service['model'],
                      messages=[{'role': 'system', 'content': 'You are a helpful assistant.'},
                                {'role': 'user',   'content': prompt}])
        cursor.setPosition(cursor.selectionEnd())
        cursor.insertText(f'\n{"-"*10}Start LLM generated\n{response.choices[0].message.content.strip()}\n{"-"*10}End LLM generated\n')
        return


    def toggleBold(self):
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if not self.editor.fontWeight() == QFont.Bold else QFont.Normal)
        self.mergeFormat(fmt)

    def toggleItalic(self):
        fmt = QTextCharFormat()
        fmt.setFontItalic(not self.editor.fontItalic())
        self.mergeFormat(fmt)

    def toggleUnderline(self):
        fmt = QTextCharFormat()
        fmt.setFontUnderline(not self.editor.fontUnderline())
        self.mergeFormat(fmt)

    def mergeFormat(self, fmt: QTextCharFormat):
        cursor = self.editor.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        self.editor.mergeCurrentCharFormat(fmt)

    def saveDocx(self):
        filename, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Files (*.docx)")
        if filename:
            doc = Document()
            html = self.editor.toHtml()
            soup = BeautifulSoup(html, "html.parser")
            for para in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                p = doc.add_paragraph()
                self._add_runs_with_formatting(p, para)
            doc.save(filename)

    def _add_runs_with_formatting(self, docx_paragraph, html_element):
        for elem in html_element.descendants:
            if isinstance(elem, str):
                docx_paragraph.add_run(elem)
            elif elem.name in ['b', 'strong']:
                run = docx_paragraph.add_run(elem.get_text())
                run.bold = True
            elif elem.name in ['i', 'em']:
                run = docx_paragraph.add_run(elem.get_text())
                run.italic = True
            elif elem.name == 'u':
                run = docx_paragraph.add_run(elem.get_text())
                run.underline = True


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = Wallo()
    win.resize(800, 600)
    win.show()
    sys.exit(app.exec())
